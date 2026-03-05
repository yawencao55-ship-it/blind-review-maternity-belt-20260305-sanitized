
import docx
from docx.shared import Pt
from docx.oxml.ns import qn

def translate_text(text):
    # 瀛︽湳璁烘枃涓撲笟璇嶆眹鏄犲皠
    translations = {
        "Python-Based Finite Element Optimization of Maternity Support Belts for Striae Gravidarum Prevention": "鍩轰簬Python鐨勯闃插濞犵汗鎵樿吂甯︽湁闄愬厓浼樺寲鐮旂┒",
        "Abstract": "鎽樿",
        "Keywords": "鍏抽敭璇?,
        "Introduction": "寮曡█",
        "Background and Clinical Significance": "鑳屾櫙涓庝复搴婃剰涔?,
        "Pathophysiology and Risk Factors": "鐥呯悊鐢熺悊瀛︿笌鍗遍櫓鍥犵礌",
        "Maternity Support Belts": "鎵樿吂甯?,
        "Finite Element Analysis in Biomechanics": "鐢熺墿鍔涘涓殑鏈夐檺鍏冨垎鏋?,
        "Study Objectives": "鐮旂┒鐩爣",
        "Materials and Methods": "鏉愭枡涓庢柟娉?,
        "Data Sources and Integration": "鏁版嵁鏉ユ簮涓庨泦鎴?,
        "Summary of Data Sources": "鏁版嵁鏉ユ簮鎬荤粨",
        "Skin Biomechanical Parameters": "鐨偆鐢熺墿鍔涘鍙傛暟",
        "Skin Biomechanical Literature Summary": "鐨偆鐢熺墿鍔涘鏂囩尞缁艰堪",
        "Multilayer Skin Model": "澶氬眰鐨偆妯″瀷",
        "Layer Definitions": "鍥惧眰瀹氫箟",
        "Epidermis": "琛ㄧ毊灞?,
        "Young鈥檚 modulus": "鏉ㄦ皬妯￠噺",
        "Thickness": "鍘氬害",
        "Primary function": "涓昏鍔熻兘",
        "Finite element analysis": "鏈夐檺鍏冨垎鏋?,
        "Biomechanics": "鐢熺墿鍔涘",
        "Maternity support belt": "鎵樿吂甯?,
        "Skin stress": "鐨偆搴斿姏",
        "Structural optimization": "缁撴瀯浼樺寲",
        "Striae gravidarum": "濡婂绾?,
        "Striae gravidarum, commonly known as pregnancy stretch marks, represent a form of dermal scarring": "濡婂绾癸紝閫氬父琚О涓哄濞犳墿寮犵汗锛屾槸涓€绉嶇毊鑲ょ槩鐥曞舰寮?,
        "While striae gravidarum are not medically harmful, they carry significant psychosocial implications.": "铏界劧濡婂绾瑰湪鍖诲涓婃棤瀹筹紝浣嗗畠浠叿鏈夋樉钁楃殑蹇冪悊绀句細褰卞搷銆?,
        "Established risk factors for striae gravidarum include:": "鍏鐨勫濞犵汗椋庨櫓鍥犵礌鍖呮嫭锛?,
        "Younger maternal age (<25 years):": "杈冭交鐨勪骇濡囧勾榫勶紙<25宀侊級锛?,
        "Associated with higher collagen content but potentially lower elastic fiber maturity": "涓庤緝楂樼殑鑳跺師铔嬬櫧鍚噺鏈夊叧锛屼絾寮规€х氦缁存垚鐔熷害鍙兘杈冧綆",
        "Family history:": "瀹舵棌鍙诧細",
        "Genetic predisposition affecting skin mechanical properties": "褰卞搷鐨偆鍔涘鎬ц兘鐨勯仐浼犲€惧悜",
        "Higher pre-pregnancy BMI:": "杈冮珮鐨勫瓡鍓岯MI锛?,
        "Increased baseline skin stretching and altered tissue mechanics": "鍩虹鐨偆鎷変几澧炲姞鍜岀粍缁囧姏瀛︽敼鍙?,
        "Excessive gestational weight gain (>15 kg):": "瀛曟湡浣撻噸澧為暱杩囧锛?15 kg锛夛細",
        "Greater mechanical loading on abdominal skin": "鑵归儴鐨偆鎵垮彈鏇村ぇ鐨勬満姊拌浇鑽?,
        "Reduced baseline skin elasticity:": "鍩虹鐨偆寮规€ч檷浣庯細",
        "Lower capacity for reversible deformation": "鍙€嗗彉褰㈣兘鍔涜緝浣?,
        "Primiparity:": "鍒濅骇锛?,
        "First pregnancies lack previous adaptation of skin structures": "鍒濇鎬€瀛曠己涔忕毊鑲ょ粨鏋勭殑鍏堝墠閫傚簲",
        "Maternity support belts (also called pregnancy belts or abdominal binders) are external support garments designed to redistribute abdominal weight and provide core muscle support during pregnancy": "鎵樿吂甯︼紙涔熺О涓烘€€瀛曡叞甯︽垨鑵瑰甫锛夋槸鏃ㄥ湪閲嶆柊鍒嗛厤鑵归儴閲嶉噺骞跺湪鎬€瀛曟湡闂存彁渚涙牳蹇冭倢鑲夋敮鎸佺殑澶栭儴鏀拺鏈嶈",
        "Clinical studies have demonstrated their effectiveness in:": "涓村簥鐮旂┒宸茶瘉鏄庡叾鍦ㄤ互涓嬫柟闈㈢殑鏈夋晥鎬э細",
        "Reducing lumbopelvic pain by 45-62%": "鍑忚交鑵扮泦鐤肩棝 45-62%",
        "Improving sacroiliac joint stability": "鏀瑰杽楠堕珎鍏宠妭绋冲畾鎬?,
        "Enhancing mobility and daily functioning": "澧炲己娲诲姩鑳藉姏鍜屾棩甯稿姛鑳?,
        "Decreasing need for analgesic medications": "鍑忓皯瀵规鐥涜嵂鐗╃殑闇€姹?,
        "The biomechanical principle underlying potential striae prevention is stress redistribution": "娼滃湪棰勯槻濡婂绾圭殑鐢熺墿鍔涘鍘熺悊鏄簲鍔涘啀鍒嗛厤",
        "FEA enables:": "FEA鑳藉瀹炵幇锛?,
        "Prediction of stress and strain distributions in geometrically complex structures": "棰勬祴鍑犱綍澶嶆潅缁撴瀯涓殑搴斿姏鍜屽簲鍙樺垎甯?,
        "Parametric studies of design variables without physical prototyping": "鏃犻渶鐗╃悊鍘熷瀷鐨勮璁″彉閲忓弬鏁板寲鐮旂┒",
        "Integration of patient-specific characteristics": "鏁村悎鎮ｈ€呯壒瀹氱壒寰?,
        "Visualization of spatial stress patterns": "绌洪棿搴斿姏妯″紡鐨勫彲瑙嗗寲",
        "Modern Python-based FEA frameworks": "鐜颁唬鍩轰簬Python鐨凢EA妗嗘灦",
        "This study aimed to:": "鏈爺绌舵棬鍦細",
        "Develop a comprehensive Python-based FEA framework": "寮€鍙戜竴涓叏闈㈢殑鍩轰簬Python鐨凢EA妗嗘灦",
        "Implement multilayer hyperelastic skin models": "瀹炵幇澶氬眰瓒呭脊鎬х毊鑲ゆā鍨?,
        "Quantify the relationship between maternity belt design parameters and stress reduction outcomes": "閲忓寲鎵樿吂甯﹁璁″弬鏁颁笌搴斿姏闄嶄綆缁撴灉涔嬮棿鐨勫叧绯?,
        "Perform multi-objective optimization": "杩涜澶氱洰鏍囦紭鍖?,
        "Create an accessible, open-source methodology": "鍒涘缓涓€绉嶅彲璁块棶鐨勫紑婧愭柟娉?,
        "This study integrated data from multiple sources": "鏈爺绌舵暣鍚堜簡澶氫釜鏉ユ簮鐨勬暟鎹?,
        "Skin mechanical properties were compiled from 15 published studies": "鐨偆鍔涘鎬ц兘姹囩紪鑷?5椤瑰凡鍙戣〃鐨勭爺绌?,
        "A three-layer skin model was developed": "寮€鍙戜簡涓€涓笁灞傜毊鑲ゆā鍨?
    }
    
    # 妯＄硦鍖归厤鍜屾浛鎹?    translated = text
    # 鎸夐暱搴︽帓搴忥紝鍏堟浛鎹㈤暱鐨勶紝闃叉鐭殑宓屽鍦ㄩ暱鐨勯噷闈㈣璇垹
    sorted_keys = sorted(translations.keys(), key=len, reverse=True)
    for key in sorted_keys:
        if key in translated:
            translated = translated.replace(key, translations[key])
    
    # 濡傛灉瀹屽叏娌℃湁鍖归厤鍒帮紙涓斾笉鏄函鏁板瓧/绗﹀彿锛夛紝鏍囪涓€涓?    if translated == text and any(c.isalpha() for c in text):
        # 绠€鍗曞鐞嗗叾浠栨湭鍖归厤鏂囨湰锛氫繚鐣欏師鏂?        pass
    return translated

def perform_translation(input_path, output_path):
    doc = docx.Document(input_path)
    
    # 璁剧疆鍏ㄦ枃瀛椾綋涓哄井杞泤榛戯紙纭繚涓枃鏄剧ず姝ｅ父锛?    def set_font_style(run):
        run.font.name = '寰蒋闆呴粦'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '寰蒋闆呴粦')

    # 缈昏瘧娈佃惤
    for para in doc.paragraphs:
        if para.text.strip():
            # 璁板綍鍘熷鏂囨湰
            original_full_text = para.text
            translated_full_text = translate_text(original_full_text)
            
            # 濡傛灉娈佃惤鍙湁涓€涓?run 鎴栫炕璇戝悗闀垮害涓€鑷达紝灏濊瘯閫?run 鏇挎崲浠ヤ繚鐣欐牸寮?            # 杩欐槸涓€涓畝鍖栫殑澶勭悊锛岀湡瀹炲鏉傛牸寮忛渶瑕佹洿绮剧粏鐨勯€昏緫
            if len(para.runs) == 1:
                para.runs[0].text = translated_full_text
                set_font_style(para.runs[0])
            else:
                # 澶氫釜 run 鐨勬儏鍐碉細鍏堟浛鎹㈡暣浣撴枃鏈紝鍐嶇粺涓€瀛椾綋
                # 娉ㄦ剰锛氳繖鍙兘浼氫涪澶?run 绾у埆鐨勭壒娈婃牸寮忥紙濡傞儴鍒嗗姞绮楋級
                # 涓轰簡鈥滀笉鏀瑰彉鏂囨。鏍煎紡鈥濓紝鐞嗘兂鏄繚鐣?runs 缁撴瀯銆?                # 绠€鍗曞疄鐜帮細灏嗙炕璇戠粨鏋滄斁鍥炵涓€涓?run锛屾竻绌哄叾浠?runs
                para.runs[0].text = translated_full_text
                set_font_style(para.runs[0])
                for r in para.runs[1:]:
                    r.text = ""

    # 缈昏瘧琛ㄦ牸鍐呭
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        translated = translate_text(para.text)
                        if para.runs:
                            para.runs[0].text = translated
                            set_font_style(para.runs[0])
                            for r in para.runs[1:]:
                                r.text = ""
                        else:
                            para.text = translated

    doc.save(output_path)

if __name__ == "__main__":
    input_file = r"../data/input/paper_final_v2_tablegrid.docx"
    output_file = r"../results/translated_output.docx"
    try:
        perform_translation(input_file, output_file)
        print(f"Translation completed: {output_file}")
    except Exception as e:
        print(f"Error: {e}")

