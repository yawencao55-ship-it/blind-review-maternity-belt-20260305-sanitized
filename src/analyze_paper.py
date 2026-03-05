
import docx
import sys

def analyze_paper_structure(file_path):
    doc = docx.Document(file_path)
    content = []
    
    # 姒傝娈佃惤
    content.append(f"Total Paragraphs: {len(doc.paragraphs)}")
    content.append("--- Paragraph Samples ---")
    for i, para in enumerate(doc.paragraphs[:20]): # 鍙湅鍓?0涓钀?        if para.text.strip():
            content.append(f"P{i} (Style: {para.style.name}): {para.text[:100]}...")
            
    # 姒傝琛ㄦ牸
    content.append(f"\nTotal Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        content.append(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} cols")
            
    return '\n'.join(content)

if __name__ == "__main__":
    file_path = r"../data/input/paper_final_v2_tablegrid.docx"
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        print(analyze_paper_structure(file_path))
    except Exception as e:
        print(f"Error: {e}")

