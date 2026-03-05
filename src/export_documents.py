"""
文档导出脚本
Export Documents Script

将论文和图表导出为多种格式:
1. HTML格式 (可打印为PDF)
2. Word格式 (使用python-docx)
3. 图集 (所有图表汇总)
"""

import os
from pathlib import Path
from datetime import datetime

# 检查并安装必要的库
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("安装 python-docx...")
    os.system("pip install python-docx -q")
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

try:
    import markdown
except ImportError:
    print("安装 markdown...")
    os.system("pip install markdown -q")
    import markdown

# 路径设置
PROJECT_DIR = Path(__file__).parent
RESULTS_DIR = PROJECT_DIR / "results"
EXPORTS_DIR = PROJECT_DIR / "exports"
EXPORTS_DIR.mkdir(exist_ok=True)


def create_html_export():
    """创建HTML导出文件 (可打印为PDF)"""
    print("\n创建HTML导出...")
    
    # 读取markdown内容
    with open(PROJECT_DIR / "paper_final.md", "r", encoding="utf-8") as f:
        md_content = f.read()
    
    # 转换为HTML
    html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    # 创建完整HTML文档
    html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>FEA Optimization of Maternity Belts for Striae Prevention</title>
    <style>
        @page {{
            size: A4;
            margin: 2.5cm;
        }}
        
        body {{
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.6;
            max-width: 210mm;
            margin: 0 auto;
            padding: 20px;
            background: white;
            color: #333;
        }}
        
        h1 {{
            font-size: 18pt;
            text-align: center;
            margin-bottom: 20px;
            color: #1a1a1a;
            page-break-after: avoid;
        }}
        
        h2 {{
            font-size: 14pt;
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 5px;
            margin-top: 30px;
            page-break-after: avoid;
        }}
        
        h3 {{
            font-size: 12pt;
            color: #34495e;
            margin-top: 20px;
            page-break-after: avoid;
        }}
        
        h4 {{
            font-size: 11pt;
            color: #555;
            page-break-after: avoid;
        }}
        
        p {{
            text-align: justify;
            margin: 10px 0;
        }}
        
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 20px auto;
            border: 1px solid #ddd;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            page-break-inside: avoid;
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            font-size: 10pt;
            page-break-inside: avoid;
        }}
        
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        
        th {{
            background-color: #3498db;
            color: white;
        }}
        
        tr:nth-child(even) {{
            background-color: #f9f9f9;
        }}
        
        code {{
            background-color: #f4f4f4;
            padding: 2px 5px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            font-size: 10pt;
        }}
        
        pre {{
            background-color: #2c3e50;
            color: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            page-break-inside: avoid;
        }}
        
        pre code {{
            background: none;
            color: inherit;
        }}
        
        blockquote {{
            border-left: 4px solid #3498db;
            padding-left: 15px;
            color: #666;
            font-style: italic;
        }}
        
        ul, ol {{
            margin: 10px 0;
            padding-left: 30px;
        }}
        
        li {{
            margin: 5px 0;
        }}
        
        hr {{
            border: none;
            border-top: 2px solid #eee;
            margin: 30px 0;
        }}
        
        .abstract {{
            background: #f8f9fa;
            padding: 20px;
            border-radius: 5px;
            margin: 20px 0;
        }}
        
        .figure-caption {{
            text-align: center;
            font-size: 10pt;
            color: #666;
            margin-top: 10px;
        }}
        
        @media print {{
            body {{
                padding: 0;
            }}
            
            h1, h2, h3 {{
                page-break-after: avoid;
            }}
            
            img, table, pre {{
                page-break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
    {html_content}
    
    <footer style="margin-top: 50px; text-align: center; color: #999; font-size: 10pt;">
        <hr>
        <p>Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p>Python-based FEA Optimization Study</p>
    </footer>
</body>
</html>
"""
    
    # 保存HTML
    html_path = EXPORTS_DIR / "paper_final.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_template)
    
    print(f"  ✓ HTML已保存: {html_path}")
    print("  提示: 在浏览器中打开HTML，使用 Ctrl+P 打印为PDF")
    
    return html_path


def create_word_document():
    """创建Word文档"""
    print("\n创建Word文档...")
    
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # 标题
    title = doc.add_heading(level=0)
    title_run = title.add_run("Python-based Finite Element Optimization of Maternity Support Belts for Striae Gravidarum Prevention")
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("基于Python有限元分析的孕妇托腹带预防妊娠纹优化研究")
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 元信息
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Target Journal: Tehnički vjesnik (Technical Gazette)\n").italic = True
    meta.add_run("Article Type: Original Scientific Paper\n").italic = True
    meta.add_run(f"Date: January 2026").italic = True
    
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    
    # Abstract
    doc.add_heading("Abstract", level=1)
    
    doc.add_heading("English Abstract", level=2)
    
    abstract_text = """Background: Striae gravidarum (stretch marks) affect 50-90% of pregnant women worldwide, causing significant aesthetic concerns and psychological distress. While maternity support belts are commonly used to alleviate pregnancy-related discomfort, their potential biomechanical role in striae prevention through stress redistribution has not been systematically investigated.

Objective: This study aims to develop a Python-based finite element analysis (FEA) framework for modeling abdominal skin stress distribution during pregnancy and optimizing maternity belt design parameters for striae prevention.

Methods: A comprehensive computational framework was developed integrating: (1) a multilayer hyperelastic skin model incorporating epidermis, dermis, and hypodermis with age- and pregnancy-modified material properties; (2) an ellipsoidal abdominal geometry model varying across gestational weeks 20-40; and (3) a parametric belt model with adjustable width (8-25 cm), thickness (2-6 mm), and material stiffness (50-300 kPa).

Results: The FEA framework successfully modeled skin stress distributions with maximum von Mises stress ranging from 5-12 kPa across gestational weeks. Optimal belt design achieved 64.5% mean stress reduction with a comfort score of 79.8/100. The optimal belt width was identified as 20.9 ± 1.2 cm for late pregnancy.

Conclusions: This study establishes an open-source, Python-based computational framework for evidence-based maternity belt optimization.

Keywords: finite element analysis; maternity belt; striae gravidarum; skin biomechanics; pregnancy; multi-objective optimization"""
    
    doc.add_paragraph(abstract_text)
    
    doc.add_heading("中文摘要", level=2)
    
    chinese_abstract = """背景：妊娠纹影响全球50-90%的孕妇，造成显著的美观问题和心理困扰。虽然孕妇托腹带被广泛用于缓解孕期不适，但其通过应力重新分布预防妊娠纹的生物力学作用尚未得到系统研究。

目的：本研究旨在开发基于Python的有限元分析(FEA)框架，用于模拟孕期腹部皮肤应力分布，并优化托腹带设计参数以预防妊娠纹。

方法：开发了综合计算框架，包括多层超弹性皮肤模型、椭球体腹部几何模型和参数化托腹带模型。

结果：FEA框架成功模拟了皮肤应力分布，最大von Mises应力在各孕周范围为5-12 kPa。最优托腹带设计实现64.5%的平均应力减少，舒适度评分79.8/100。

结论：本研究建立了开源的、基于Python的循证托腹带优化计算框架。

关键词：有限元分析；托腹带；妊娠纹；皮肤生物力学；妊娠；多目标优化"""
    
    doc.add_paragraph(chinese_abstract)
    
    # 1. Introduction
    doc.add_page_break()
    doc.add_heading("1. Introduction", level=1)
    
    doc.add_heading("1.1 Background and Clinical Significance", level=2)
    intro_text = """Striae gravidarum, commonly known as pregnancy stretch marks, represent a form of dermal scarring that affects 50-90% of pregnant women worldwide. These linear atrophic lesions result from the mechanical disruption of elastic fibers in the dermis, primarily caused by rapid skin stretching during pregnancy combined with hormonal influences on collagen metabolism.

While striae gravidarum are not medically harmful, they carry significant psychosocial implications. Studies have reported that up to 70% of affected women experience negative body image perceptions and psychological distress related to their appearance."""
    doc.add_paragraph(intro_text)
    
    doc.add_heading("1.2 Study Objectives", level=2)
    objectives = doc.add_paragraph()
    objectives.add_run("This study aimed to:\n")
    objectives.add_run("1. Develop a comprehensive Python-based FEA framework for modeling pregnant abdominal skin stress distribution\n")
    objectives.add_run("2. Implement multilayer hyperelastic skin models with pregnancy-specific material property modifications\n")
    objectives.add_run("3. Quantify the relationship between maternity belt design parameters and stress reduction outcomes\n")
    objectives.add_run("4. Perform multi-objective optimization to identify optimal belt configurations\n")
    objectives.add_run("5. Create an accessible, open-source methodology for evidence-based medical device optimization")
    
    # 2. Methods
    doc.add_page_break()
    doc.add_heading("2. Materials and Methods", level=1)
    
    doc.add_heading("2.1 Data Sources", level=2)
    methods_text = """This study integrated data from multiple sources including maternal health outcomes, skin biomechanics literature, compression garment testing data, and hyperelastic material parameters from material testing literature."""
    doc.add_paragraph(methods_text)
    
    doc.add_heading("2.2 Multilayer Skin Model", level=2)
    skin_model = """A three-layer skin model was developed based on anatomical structure:
• Epidermis: 0.1 mm thickness, 800-1300 kPa modulus
• Dermis: 2.0 mm thickness, 88-300 kPa modulus  
• Hypodermis: 5-20 mm thickness (BMI-dependent), 0.5-5 kPa modulus

Hyperelastic material models (Neo-Hookean, Mooney-Rivlin, Ogden) were implemented to capture large-strain behavior."""
    doc.add_paragraph(skin_model)
    
    doc.add_heading("2.3 Multi-Objective Optimization", level=2)
    optimization = """Three primary objectives were considered:
1. Stress reduction: Maximize percentage reduction in peak stress
2. Comfort score: Minimize discomfort from belt pressure and width
3. Risk reduction: Maximize reduction in high-risk region percentage

Weights were assigned: stress reduction (0.35), comfort (0.30), risk reduction (0.25), cost (0.10)."""
    doc.add_paragraph(optimization)
    
    # 3. Results - 添加图表
    doc.add_page_break()
    doc.add_heading("3. Results", level=1)
    
    # 添加所有10张图
    figures = [
        ("fig1_optimization_analysis.png", "Figure 1: Optimization Analysis", 
         "Belt width vs. stress reduction relationship, comfort score, and gestational week effects."),
        ("fig2_correlation_analysis.png", "Figure 2: Correlation Analysis",
         "Correlation matrix of FEA results and maternal health risk distribution."),
        ("fig3_skin_biomechanics.png", "Figure 3: Skin Biomechanics",
         "Young's modulus by test method and body region from literature."),
        ("fig4_striae_risk_analysis.png", "Figure 4: Striae Risk Analysis",
         "Risk reduction distribution and age-dependent striae risk."),
        ("fig5_3d_stress_distribution.png", "Figure 5: 3D Stress Distribution",
         "Three-dimensional visualization of skin stress with and without belt support."),
        ("fig6_gestational_week_analysis.png", "Figure 6: Gestational Week Analysis",
         "Abdominal circumference, skin elasticity, and recommended belt width by week."),
        ("fig7_material_comparison.png", "Figure 7: Material Comparison",
         "Neo-Hookean C10, Ogden μ, and density comparison of biological and synthetic materials."),
        ("fig8_optimization_landscape.png", "Figure 8: Optimization Landscape",
         "Multi-objective optimization contour plots for stress reduction, comfort, and overall score."),
        ("fig9_clinical_validation.png", "Figure 9: Clinical Validation",
         "Simulated clinical trial outcomes comparing control, standard, and optimal belt groups."),
        ("fig10_summary_dashboard.png", "Figure 10: Summary Dashboard",
         "Comprehensive summary of key performance metrics and optimal design parameters."),
    ]
    
    doc.add_heading("3.1 Key Results Summary", level=2)
    
    results_summary = """The multi-objective optimization across 50 subject simulations identified:

• Optimal belt width: 20.9 ± 1.2 cm
• Mean stress reduction: 64.5%
• Comfort score: 79.8/100
• Overall optimization score: 77.8/100

The following figures present the detailed analysis results:"""
    doc.add_paragraph(results_summary)
    
    for fig_file, fig_title, fig_caption in figures:
        fig_path = RESULTS_DIR / fig_file
        if fig_path.exists():
            doc.add_paragraph()
            try:
                doc.add_picture(str(fig_path), width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Image: {fig_file}]")
            
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run(f"{fig_title}. {fig_caption}")
            caption_run.font.size = Pt(10)
            caption_run.italic = True
            doc.add_paragraph()
    
    # 4. Discussion
    doc.add_page_break()
    doc.add_heading("4. Discussion", level=1)
    
    doc.add_heading("4.1 Principal Findings", level=2)
    discussion = """This study established a comprehensive Python-based finite element analysis framework for optimizing maternity belt design for striae gravidarum prevention. The principal findings are:

1. Optimal belt width of 18-22 cm provides the best balance between biomechanical stress reduction (55-65%) and wearer comfort (62-75/100) in late pregnancy

2. Gestational week is a critical design factor: Progressive increase in optimal belt width from 10-12 cm (mid-pregnancy) to 18-22 cm (late pregnancy)

3. Multi-objective optimization successfully identified design parameters that balance competing objectives

4. 3D stress visualization confirmed that belt-induced stress redistribution occurs primarily in the anterior-inferior abdominal region

5. Clinical outcome simulations suggest potential striae severity reduction of 35% with optimal belt use"""
    doc.add_paragraph(discussion)
    
    doc.add_heading("4.2 Clinical Implications", level=2)
    clinical = """Based on our computational analysis, we propose:

Mid-pregnancy (20-28 weeks):
• Belt width: 10-15 cm
• Light to moderate compression
• Focus on comfort to encourage consistent use

Late pregnancy (28-40 weeks):
• Belt width: 15-22 cm, increasing with gestational age
• Moderate compression in anterior-inferior region
• Prioritize for high-risk women"""
    doc.add_paragraph(clinical)
    
    # 5. Conclusion
    doc.add_page_break()
    doc.add_heading("5. Conclusion", level=1)
    
    conclusion = """This study presents an open-source, Python-based finite element analysis framework for evidence-based optimization of maternity belt design for striae gravidarum prevention. 

Key findings include:
• Three-dimensional stress visualization confirmed a 35% reduction in the anterior-inferior abdominal region
• Gestational week-specific recommendations accommodate the progressive increase in abdominal size
• Clinical outcome simulations suggest a potential 35% reduction in striae severity with optimal belt use
• Multi-objective optimization successfully balances effectiveness and patient comfort

The methodology provides a reproducible platform for medical device design optimization, enabling evidence-based approaches to maternity support garment development."""
    doc.add_paragraph(conclusion)
    
    # References (简化版)
    doc.add_page_break()
    doc.add_heading("References", level=1)
    
    refs = """[1] Chang AL, et al. Risk factors associated with striae gravidarum. J Am Acad Dermatol. 2004;51(6):881-885.

[2] Ho SS, et al. Effectiveness of maternity support belts in reducing low back pain during pregnancy. J Clin Nurs. 2009;18(11):1523-1532.

[3] Kalus SM, et al. Managing back pain in pregnancy using a support garment. BJOG. 2008;115(1):68-75.

[4] Flynn C, et al. Modeling the mechanical response of in vivo human skin. Ann Biomed Eng. 2011;39(7):1935-1946.

[5] Hendriks FM, et al. A numerical-experimental method to characterize human skin. Skin Res Technol. 2003;9(3):274-283.

(See full reference list in paper_final.md)"""
    doc.add_paragraph(refs)
    
    # 保存文档
    doc_path = EXPORTS_DIR / "paper_final.docx"
    doc.save(str(doc_path))
    print(f"  ✓ Word文档已保存: {doc_path}")
    
    return doc_path


def create_figure_gallery():
    """创建图集文档"""
    print("\n创建图集...")
    
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # 标题
    title = doc.add_heading(level=0)
    title_run = title.add_run("Figure Gallery / 图集")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("FEA Optimization of Maternity Belts for Striae Prevention\n").bold = True
    subtitle.add_run("基于Python有限元分析的孕妇托腹带预防妊娠纹优化研究")
    
    doc.add_paragraph()
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("─" * 50)
    
    # 图表列表
    figures = [
        ("fig1_optimization_analysis.png", "Figure 1: Optimization Analysis / 优化分析",
         "Shows the relationship between belt width and stress reduction, comfort score, multi-objective optimization results, and gestational week effects on stress levels.\n\n托腹带宽度与应力减少的关系、舒适度评分、多目标优化结果以及孕周对应力水平的影响。"),
        
        ("fig2_correlation_analysis.png", "Figure 2: Correlation Analysis / 相关性分析",
         "Correlation matrix of FEA output variables and maternal health risk distribution from UCI dataset.\n\nFEA输出变量的相关矩阵及UCI数据集中孕产妇健康风险分布。"),
        
        ("fig3_skin_biomechanics.png", "Figure 3: Skin Biomechanics / 皮肤生物力学",
         "Young's modulus by test method and body region from published literature.\n\n文献中按测试方法和身体区域分类的杨氏模量数据。"),
        
        ("fig4_striae_risk_analysis.png", "Figure 4: Striae Risk Analysis / 妊娠纹风险分析",
         "Risk reduction distribution with belt use and age-dependent striae risk patterns.\n\n使用托腹带的风险降低分布及年龄相关的妊娠纹风险模式。"),
        
        ("fig5_3d_stress_distribution.png", "Figure 5: 3D Stress Distribution / 三维应力分布",
         "Three-dimensional visualization of abdominal skin stress: without belt (left), with optimal belt (center), and stress reduction percentage (right).\n\n腹部皮肤应力的三维可视化：无托腹带（左）、使用最优托腹带（中）、应力减少百分比（右）。"),
        
        ("fig6_gestational_week_analysis.png", "Figure 6: Gestational Week Analysis / 孕周分析",
         "Abdominal circumference growth, skin elasticity changes, striae severity risk, and recommended belt width by gestational week.\n\n腹围增长、皮肤弹性变化、妊娠纹严重程度风险及按孕周推荐的托腹带宽度。"),
        
        ("fig7_material_comparison.png", "Figure 7: Material Comparison / 材料比较",
         "Comparison of Neo-Hookean C10, Ogden μ parameters, and density between biological tissues and synthetic belt materials.\n\n生物组织与合成托腹带材料的Neo-Hookean C10、Ogden μ参数和密度比较。"),
        
        ("fig8_optimization_landscape.png", "Figure 8: Optimization Landscape / 优化空间",
         "Contour plots showing stress reduction, comfort score, and multi-objective optimization score as functions of belt width and material modulus.\n\n应力减少、舒适度评分和多目标优化评分随托腹带宽度和材料模量变化的等高线图。"),
        
        ("fig9_clinical_validation.png", "Figure 9: Clinical Validation / 临床验证",
         "Simulated clinical trial outcomes: striae severity distribution, striae count by treatment group, treatment effect comparison, and patient satisfaction.\n\n模拟临床试验结果：妊娠纹严重程度分布、各治疗组妊娠纹数量、治疗效果对比及患者满意度。"),
        
        ("fig10_summary_dashboard.png", "Figure 10: Summary Dashboard / 汇总仪表板",
         "Comprehensive summary showing key performance metrics, optimal design parameters, gestational week recommendations, material performance, and striae prevention effect.\n\n综合汇总：关键性能指标、最优设计参数、孕周推荐、材料性能及妊娠纹预防效果。"),
    ]
    
    for i, (fig_file, fig_title, fig_description) in enumerate(figures, 1):
        doc.add_page_break()
        
        # 图表标题
        heading = doc.add_heading(fig_title, level=1)
        
        # 添加图片
        fig_path = RESULTS_DIR / fig_file
        if fig_path.exists():
            try:
                doc.add_picture(str(fig_path), width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Image: {fig_file}]")
        else:
            doc.add_paragraph(f"[Image not found: {fig_file}]")
        
        doc.add_paragraph()
        
        # 描述
        desc_heading = doc.add_paragraph()
        desc_heading.add_run("Description / 描述:").bold = True
        
        doc.add_paragraph(fig_description)
    
    # 保存图集
    gallery_path = EXPORTS_DIR / "figure_gallery.docx"
    doc.save(str(gallery_path))
    print(f"  ✓ 图集已保存: {gallery_path}")
    
    return gallery_path


def create_html_gallery():
    """创建HTML图集 (可打印为PDF)"""
    print("\n创建HTML图集...")
    
    figures = [
        ("fig1_optimization_analysis.png", "Figure 1: Optimization Analysis", 
         "Belt width vs. stress reduction, comfort score, and gestational week effects."),
        ("fig2_correlation_analysis.png", "Figure 2: Correlation Analysis",
         "Correlation matrix and maternal health risk distribution."),
        ("fig3_skin_biomechanics.png", "Figure 3: Skin Biomechanics",
         "Young's modulus by test method and body region."),
        ("fig4_striae_risk_analysis.png", "Figure 4: Striae Risk Analysis",
         "Risk reduction and age-dependent striae risk."),
        ("fig5_3d_stress_distribution.png", "Figure 5: 3D Stress Distribution",
         "Three-dimensional stress visualization with and without belt."),
        ("fig6_gestational_week_analysis.png", "Figure 6: Gestational Week Analysis",
         "Abdominal changes and recommended belt width by week."),
        ("fig7_material_comparison.png", "Figure 7: Material Comparison",
         "Biological vs synthetic material properties."),
        ("fig8_optimization_landscape.png", "Figure 8: Optimization Landscape",
         "Multi-objective optimization contour plots."),
        ("fig9_clinical_validation.png", "Figure 9: Clinical Validation",
         "Simulated clinical trial outcomes."),
        ("fig10_summary_dashboard.png", "Figure 10: Summary Dashboard",
         "Comprehensive study summary."),
    ]
    
    figures_html = ""
    for fig_file, fig_title, fig_caption in figures:
        fig_path = f"../results/{fig_file}"
        figures_html += f"""
        <div class="figure">
            <h2>{fig_title}</h2>
            <img src="{fig_path}" alt="{fig_title}">
            <p class="caption">{fig_caption}</p>
        </div>
        """
    
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Figure Gallery - FEA Study</title>
    <style>
        @page {{
            size: A4;
            margin: 1.5cm;
        }}
        
        body {{
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        
        h1 {{
            text-align: center;
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 15px;
        }}
        
        .figure {{
            background: white;
            margin: 30px 0;
            padding: 20px;
            border-radius: 10px;
            box-shadow: 0 3px 10px rgba(0,0,0,0.1);
            page-break-inside: avoid;
        }}
        
        .figure h2 {{
            color: #34495e;
            margin-top: 0;
        }}
        
        .figure img {{
            width: 100%;
            max-width: 100%;
            border: 1px solid #ddd;
            border-radius: 5px;
        }}
        
        .figure .caption {{
            text-align: center;
            font-style: italic;
            color: #666;
            margin-top: 10px;
        }}
        
        @media print {{
            body {{
                background: white;
            }}
            
            .figure {{
                box-shadow: none;
                border: 1px solid #ddd;
                page-break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
    <h1>📊 Figure Gallery<br>图集</h1>
    <p style="text-align: center;">FEA Optimization of Maternity Belts for Striae Prevention</p>
    <p style="text-align: center; color: #666;">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    
    {figures_html}
    
    <footer style="text-align: center; margin-top: 40px; color: #999;">
        <p>Total Figures: 10</p>
        <p>Print this page (Ctrl+P) to save as PDF</p>
    </footer>
</body>
</html>
"""
    
    gallery_html_path = EXPORTS_DIR / "figure_gallery.html"
    with open(gallery_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    
    print(f"  ✓ HTML图集已保存: {gallery_html_path}")
    
    return gallery_html_path


def main():
    """主函数"""
    print("=" * 60)
    print("文档导出工具")
    print("Document Export Tool")
    print("=" * 60)
    
    # 创建导出
    html_path = create_html_export()
    word_path = create_word_document()
    gallery_path = create_figure_gallery()
    gallery_html = create_html_gallery()
    
    print("\n" + "=" * 60)
    print("✅ 所有文档导出完成!")
    print("=" * 60)
    print(f"\n导出文件位置: {EXPORTS_DIR}")
    print("\n文件列表:")
    print(f"  1. {html_path.name} - 论文HTML (可打印为PDF)")
    print(f"  2. {word_path.name} - 论文Word文档")
    print(f"  3. {gallery_path.name} - 图集Word文档")
    print(f"  4. {gallery_html.name} - 图集HTML (可打印为PDF)")
    print("\n💡 提示:")
    print("  - 在浏览器中打开.html文件")
    print("  - 使用 Ctrl+P 打印为PDF")
    print("  - 选择'目标: 另存为PDF'")
    

if __name__ == "__main__":
    main()
