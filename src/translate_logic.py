
import docx
import sys
import copy

def translate_placeholder(text):
    # 杩欓噷鏄炕璇戦€昏緫鐨勫崰浣嶇锛屽疄闄呮搷浣滀腑鎴戜細鏍规嵁鏂囨湰鍐呭杩涜楂樿川閲忕炕璇?    # 妯℃嫙缈昏瘧杩囩▼锛屽疄闄呮墽琛屾椂浼氭浛鎹负鐪熷疄鐨勭炕璇戠粨鏋?    mapping = {
        "Python-Based Finite Element Optimization of Maternity Support Belts for Striae Gravidarum Prevention": "鍩轰簬Python鐨勯闃插濞犵汗鎵樿吂甯︽湁闄愬厓浼樺寲鐮旂┒",
        "Introduction": "寮曡█",
        "Abstract": "鎽樿",
        "Keywords": "鍏抽敭璇?,
        "Finite element analysis": "鏈夐檺鍏冨垎鏋?,
        "Biomechanics": "鐢熺墿鍔涘",
        "Maternity support belt": "鎵樿吂甯?,
        "Skin stress": "鐨偆搴斿姏",
        "Structural optimization": "缁撴瀯浼樺寲",
        "Background and Clinical Significance": "鑳屾櫙涓庝复搴婃剰涔?,
        "Pathophysiology and Risk Factors": "鐥呯悊鐢熺悊瀛︿笌鍗遍櫓鍥犵礌",
        "Striae gravidarum, commonly known as pregnancy stretch marks, represent a form of dermal scarring": "濡婂绾癸紝閫氬父琚О涓哄濞犳墿寮犵汗锛屾槸涓€绉嶇毊鑲ょ槩鐥曞舰寮?,
        "While striae gravidarum are not medically harmful, they carry significant psychosocial implications.": "铏界劧濡婂绾瑰湪鍖诲涓婃棤瀹筹紝浣嗗畠浠叿鏈夋樉钁楃殑蹇冪悊绀句細褰卞搷銆?
    }
    for eng, chi in mapping.items():
        if eng in text:
            return text.replace(eng, chi)
    return f"[璇慮: {text}" # 榛樿鏍囪

def translate_docx(input_path, output_path):
    doc = docx.Document(input_path)
    
    # 缈昏瘧娈佃惤锛屽悓鏃朵繚鎸佹牱寮?    for para in doc.paragraphs:
        if para.text.strip():
            # 璁板綍鍘熷 runs 鐨勬牱寮忓苟閫愪釜鏇挎崲鏂囨湰
            original_text = para.text
            # 绠€鍗曟浛鎹㈢瓥鐣ワ細淇濇寔 paragraph 绾у埆鏍峰紡锛屾浛鎹㈡枃鏈唴瀹?            # 娉ㄦ剰锛氬鏉傛枃妗ｄ腑 run 绾у埆鏍峰紡锛堝鍔犵矖銆佹枩浣擄級闇€瑕佺簿缁嗗鐞?            translated_text = translate_placeholder(original_text)
            
            # 涓轰簡淇濇寔鏍煎紡锛屾垜浠竻绌哄師鏂囨湰浣嗕繚鐣欐牱寮?            # 姣旇緝绋冲Ε鐨勬柟娉曟槸閫愪釜 run 缈昏瘧锛屼絾 OCR/鏈哄櫒缈昏瘧閫氬父浠ュ彞瀛愪负鍗曚綅
            # 杩欓噷鍏堟紨绀烘暣浣撴枃鏈浛鎹互淇濇寔鍩烘湰鏍峰紡
            if len(para.runs) > 0:
                para.text = translated_text
                
    # 缈昏瘧琛ㄦ牸
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        paragraph.text = translate_placeholder(paragraph.text)
                        
    doc.save(output_path)

if __name__ == "__main__":
    input_file = r"../data/input/paper_final_v2_tablegrid.docx"
    output_file = r"../results/translated_output.docx"
    try:
        # 杩欓噷鍙槸婕旂ず閫昏緫锛屽疄闄呯炕璇戦渶瑕佸ぇ瑙勬ā澶勭悊
        # translate_docx(input_file, output_file)
        print("Ready for translation processing.")
    except Exception as e:
        print(f"Error: {e}")

