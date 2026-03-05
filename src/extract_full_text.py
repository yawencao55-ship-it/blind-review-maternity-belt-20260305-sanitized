
import docx
import sys

def full_text_extract(file_path):
    doc = docx.Document(file_path)
    text_blocks = []
    
    # 鎻愬彇娈佃惤
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            text_blocks.append(f"P_{i}: {para.text}")
            
    # 鎻愬彇琛ㄦ牸
    for i, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    text_blocks.append(f"T_{i}_R{r_idx}_C{c_idx}: {cell.text}")
                    
    return text_blocks

if __name__ == "__main__":
    file_path = r"../data/input/paper_final_v2_tablegrid.docx"
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        blocks = full_text_extract(file_path)
        # 杈撳嚭鍓?0涓潡浠ヤ緵缈昏瘧鍙傝€?        for block in blocks[:50]:
            print(block)
    except Exception as e:
        print(f"Error: {e}")

